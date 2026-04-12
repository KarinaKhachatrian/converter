import string

from docx import Document

from src.pdf_processor.headers import h2_headers, h3_headers
from src.pdf_processor.processor import Processor


class DocumentProcessor(Processor):
    def __init__(self, docx_path: str, counter=1):
        self.docx_path = docx_path
        self.counter = counter
        self.doc = Document(self.docx_path)

    @staticmethod
    def apply_bold(paragraph):
        for run in paragraph.runs:
            run.bold = True

    @staticmethod
    def remove_nums(numpr_elements):
        for numpr in numpr_elements:
            parent = numpr.getparent()
            parent.remove(numpr)

    @staticmethod
    def clean_text(text) -> str:
        clean_text = text.lstrip(string.punctuation).rstrip(string.punctuation)
        return clean_text.strip()

    def clean_first_paragraphs(self):
        paragraphs = self.doc.paragraphs
        for i in [0, 1]:

            paragraphs[i].text = ''
        self.doc.save(self.docx_path)

    def process(self):
        self.clean_first_paragraphs()

        for paragraph in self.doc.paragraphs:
            text = paragraph.text
            clean_text = self.clean_text(text)

            numpr_elements = paragraph._p.xpath('.//w:numPr')

            h2_check = any(clean_text.startswith(h2) for h2 in h2_headers)

            h3_check = None
            for h3_header in h3_headers.keys():
                if clean_text.startswith(h3_header):
                    h3_check = h3_header
                    break

            if h2_check and numpr_elements:
                text = f'{self.counter}. {clean_text}'
                paragraph.text = text
                self.counter += 1

                self.apply_bold(paragraph)
                self.remove_nums(numpr_elements)

            elif h3_check and numpr_elements:
                h3_num = h3_headers[h3_check]
                text = f'{h3_num} {clean_text}'
                paragraph.text = text

                self.apply_bold(paragraph)
                self.remove_nums(numpr_elements)


        self.doc.save(self.docx_path)
